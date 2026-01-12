"""
Document loading utilities for Q&A generation.
"""
import os
import re
from pathlib import Path
from typing import Tuple, Optional
import PyPDF2
from docx import Document


def extract_year_from_filename(filename: str) -> Optional[str]:
    """Extract year from filename if present."""
    year_match = re.search(r'\b(20\d{2})\b', filename)
    return year_match.group(1) if year_match else None


def load_pdf(file_path: str) -> Tuple[str, dict]:
    """
    Load PDF and extract text with page tracking.
    
    Returns:
        (text, metadata)
    """
    text = ""
    page_texts = {}
    
    try:
        with open(file_path, 'rb') as f:
            reader = PyPDF2.PdfReader(f)
            page_count = len(reader.pages)
            
            for page_num in range(page_count):
                page = reader.pages[page_num]
                page_text = page.extract_text()
                page_texts[page_num + 1] = page_text
                text += f"\n[PAGE {page_num + 1}]\n{page_text}\n"
        
        metadata = {
            'page_count': page_count,
            'page_texts': page_texts
        }
        
        return text, metadata
        
    except Exception as e:
        print(f"✗ Error loading PDF: {e}")
        return "", {}


def load_docx(file_path: str) -> Tuple[str, dict]:
    """
    Load DOCX and extract text.
    
    Returns:
        (text, metadata)
    """
    try:
        doc = Document(file_path)
        text = "\n".join([para.text for para in doc.paragraphs])
        
        metadata = {
            'paragraph_count': len(doc.paragraphs)
        }
        
        return text, metadata
        
    except Exception as e:
        print(f"✗ Error loading DOCX: {e}")
        return "", {}


def load_document(file_path: str) -> Tuple[str, str, Optional[str], dict]:
    """
    Load document and return text, title, year, metadata.
    
    Returns:
        (text, title, year, metadata)
    """
    
    file_path = Path(file_path)
    
    # Extract title and year from filename
    title = file_path.stem  # Filename without extension
    year = extract_year_from_filename(title)
    
    # Load based on extension
    ext = file_path.suffix.lower()
    
    if ext == '.pdf':
        text, metadata = load_pdf(str(file_path))
    elif ext in ['.docx', '.doc']:
        text, metadata = load_docx(str(file_path))
    elif ext in ['.txt', '.md']:
        with open(file_path, 'r', encoding='utf-8') as f:
            text = f.read()
        metadata = {}
    else:
        print(f"✗ Unsupported file format: {ext}")
        return "", title, year, {}
    
    metadata['filename'] = file_path.name
    metadata['extension'] = ext
    
    return text, title, year, metadata

